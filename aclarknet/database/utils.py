from boto.exception import BotoServerError
from collections import OrderedDict
from decimal import Decimal
from django.conf import settings as django_settings
from django.contrib.gis.geoip2 import GeoIP2
from django.contrib import messages
from django.core.mail import send_mail
from django.core.paginator import Paginator
from django.core.paginator import EmptyPage
from django.core.paginator import PageNotAnInteger
from django.core.urlresolvers import reverse
from django.db.models import Q
from django.db.models import F
from django.db.models import Sum
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.shortcuts import get_object_or_404
from django.shortcuts import render
from django.template.loader import render_to_string
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from faker import Faker
from functools import reduce
from import_export import widgets
from hashlib import md5
from io import BytesIO
from io import StringIO
from lxml import etree
from matplotlib.dates import DateFormatter
from matplotlib.dates import date2num
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg
from operator import or_ as OR

fake = Faker()
geo_ip = GeoIP2()

URL_NAMES = {
    'Settings App': ('settings_app', 'settings_app_edit', ''),  # custom meta
    'Settings Company': ('settings_company', 'settings_company_edit', ''),
    'Settings Contract': ('settings_contract', 'settings_contract_edit', ''),
    'client': ('client_view', 'client_edit', 'client_index'),
    'contact': ('contact_view', 'contact_edit', 'contact_index'),
    'contract': ('contract_view', 'contract_edit', 'contract_index'),
    'estimate': ('estimate_view', 'estimate_edit', 'estimate_index'),
    'file': ('file_view', 'file_edit', 'file_index'),
    'invoice': ('invoice_view', 'invoice_edit', 'invoice_index'),
    'log': ('log_view', 'log_edit', 'log_index'),
    'newsletter': ('newsletter_view', 'newsletter_edit', 'newsletter_index'),
    'note': ('note_view', 'note_edit', 'note_index'),
    'profile': ('user_view', 'user_edit', 'user_index'),
    'project': ('project_view', 'project_edit', 'project_index'),
    'proposal': ('proposal_view', 'proposal_edit', 'proposal_index'),
    'report': ('report_view', 'report_edit', 'report_index'),
    'service': ('', 'service_edit', ''),
    'task': ('task_view', 'task_edit', 'task_index'),
    'time': ('time_view', 'time_edit', 'time_index'),
    'user': ('user_view', 'user_edit', 'user_index'),
}


class BooleanWidget(widgets.Widget):
    """
    Convert strings to boolean values
    """

    def clean(self, value):
        if value == 'Yes':
            return True
        else:
            return False


class DecimalWidget(widgets.Widget):
    """
    Convert strings to decimal values
    """

    def clean(self, value):
        if value:
            return Decimal(value.replace(',', ''))
        else:
            return Decimal(0)


class UserWidget(widgets.Widget):
    """
    """

    def clean(self, value):
        return value


def edit(request, **kwargs):
    context = {}
    obj = None
    app_settings_model = kwargs.get('app_settings_model')
    client_model = kwargs.get('client_model')
    company_model = kwargs.get('company_model')
    contact_model = kwargs.get('contact_model')
    estimate_model = kwargs.get('estimate_model')
    form_model = kwargs.get('form_model')
    invoice_model = kwargs.get('invoice_model')
    model = kwargs.get('model')
    note_model = kwargs.get('note_model')
    pk = kwargs.get('pk')
    project_model = kwargs.get('project_model')
    user_model = kwargs.get('user_model')
    model_name = None
    if model:
        model_name = model._meta.verbose_name
        context['active_nav'] = model_name
    if pk is None:  # New or mail
        form = get_form(
            client_model=client_model,
            form_model=form_model,
            invoice_model=invoice_model,
            model=model,
            user_model=user_model,
            request=request)
    else:  # Existing
        obj = get_object_or_404(model, pk=pk)
        if model_name == 'user':  # One-off to edit user profile
            obj = obj.profile
        form = get_form(form_model=form_model, obj=obj)
    if request.method == 'POST':
        ref = request.META.get('HTTP_REFERER')
        if pk is None:
            if model_name == 'user':  # One-off to create user
                username = fake.text()[:150]
                new_user = model.objects.create_user(username=username)
            form = form_model(request.POST)
        else:
            copy = request.POST.get('copy')  # Copy or delete
            delete = request.POST.get('delete')
            if copy:
                return obj_copy(obj)
            if delete:
                return obj_remove(obj)
            query_checkbox = get_query(request, 'checkbox')  # Check boxes
            if query_checkbox['condition']:
                return set_check_boxes(obj, query_checkbox, ref,
                                       app_settings_model)
            sent = request.POST.get('sent')  # Invoice sent
            if sent:
                return obj_sent(obj, ref)
            form = form_model(request.POST, instance=obj)
        if form.is_valid():
            try:  # New object
                obj = form.save()
                if model_name == 'user':  # One-off to create profile
                    if not obj.user:  # for new user
                        obj.user = new_user
                        obj.save()
                set_relationship(
                    obj,
                    request,
                    client_model=client_model,
                    company_model=company_model,
                    estimate_model=estimate_model,
                    invoice_model=invoice_model,
                    model=model,
                    project_model=project_model)
                if model_name == 'time':
                    status_message = {
                        'success': 'Time entry created',
                        'failure': 'Time entry not created',
                    }
                    mail_process(
                        obj,
                        form=form,
                        status_message=status_message,
                        request=request)
                return obj_redir(obj, pk=pk)
            except AttributeError:  # No new object. Just sending mail.
                status_message = {
                    'success': 'Mail sent to %s!',
                    'failure': 'Mail not sent to %s!',
                }
                obj = mail_obj(
                    request,
                    contact_model=contact_model,
                    estimate_model=estimate_model,
                    note_model=note_model)
                mail_process(
                    obj,
                    form=form,
                    status_message=status_message,
                    request=request)
    context['form'] = form
    context['is_staff'] = request.user.is_staff
    context['item'] = obj
    context['pk'] = pk
    if company_model:
        company = company_model.get_solo()
        context['company'] = company
    if invoice_model:  # Dashboard totals for reporting
        invoices = invoice_model.objects.filter(last_payment_date=None)
        gross, net = get_invoice_totals(invoices)
        context['gross'] = gross
        context['net'] = net
    elif contact_model:
        model_name = contact_model._meta.verbose_name
    elif note_model:
        model_name = note_model._meta.verbose_name
    template_name = get_template_and_url(
        model_name=model_name, page_type='edit')
    return render(request, template_name, context)


def generate_doc(contract):
    """
    https://stackoverflow.com/a/24122313/185820
    """
    document = Document()
    # Head
    task = ''
    if contract.task:
        task = contract.task
    title = document.add_heading(
        'ACLARK.NET, LLC %s AGREEMENT PREPARED FOR:' % task, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if contract.client:
        client_name = document.add_heading(contract.client.name, level=1)
        client_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_address = document.add_heading(contract.client.address, level=1)
        client_address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parser = etree.HTMLParser()  # http://lxml.de/parsing.html
    tree = etree.parse(StringIO(contract.body), parser)
    # Body
    for element in tree.iter():
        if element.tag == 'h2':
            document.add_heading(element.text, level=2)
        elif element.tag == 'p':
            document.add_paragraph(element.text)
    return document


def get_geo_ip_data(request, ip_address=None):
    if not ip_address:
        # https://stackoverflow.com/a/4581997/185820
        ip_address = request.META.get('HTTP_X_REAL_IP')
    if ip_address:
        return geo_ip.city(ip_address)


def get_company_name(model):
    company = model.get_solo()
    if company.name:
        company_name = company.name
    else:
        company_name = fake.text()
    # Ghetto
    company_name = company.name.replace('.', '_')
    company_name = company_name.replace(', ', '_')
    company_name = company_name.replace('#', '_')
    company_name = company_name.replace('-', '_')
    company_name = company_name.replace('(', '_')
    company_name = company_name.replace(')', '_')
    company_name = company_name.replace(' ', '_')
    company_name = company_name.upper()
    return company_name


def get_fields(items, exclude_fields=[]):
    for item in items:
        fields = item._meta._get_fields()
        item.fields = OrderedDict()
        for field in fields:
            if not field.is_relation and field.name not in exclude_fields:
                field_name = field.name.title().replace('_', ' ')
                value = getattr(item, field.name)
                if value:
                    try:
                        value = value.title()
                    except:  # Probably not "regular" field
                        pass
                item.fields[field_name] = value
    return items


def get_form(**kwargs):
    """
    Return appropriate form based on new or edit
    """
    client_model = kwargs.get('client_model')
    form_model = kwargs.get('form_model')
    invoice_model = kwargs.get('invoice_model')
    model = kwargs.get('model')
    obj = kwargs.get('obj')
    request = kwargs.get('request')
    user_model = kwargs.get('user_model')
    query_client = None
    query_user = None
    if request:
        query_user = get_query(request, 'user')
        query_client = get_query(request, 'client')
    if obj:  # Existing object
        model_name = obj._meta.verbose_name
        if model_name == 'note':  # Populate form with tags already set
            form = form_model(initial={'tags': obj.tags.all()}, instance=obj)
        else:
            form = form_model(instance=obj)
    else:  # New object or mail
        initial = {'send_html': True}
        if model:
            model_name = model._meta.verbose_name
            if model_name == 'report' and invoice_model:  # Populate new report
                # with gross, net.
                invoices = invoice_model.objects.filter(last_payment_date=None)
                gross, net = get_invoice_totals(invoices)
                obj = model(gross=gross, net=net)
                form = form_model(instance=obj)
            elif model_name == 'contact':  # Populate new contact
                # with appropriate fields
                if query_user:
                    user = get_object_or_404(user_model, pk=query_user)
                    obj = model(email=user.email)
                elif query_client:
                    client = get_object_or_404(client_model, pk=query_client)
                    obj = model(client=client)
                form = form_model(instance=obj)
            elif model_name == 'estimate':
                user = get_object_or_404(user_model, pk=query_user)
                obj = model(user=user, is_to=True)
                form = form_model(instance=obj)
            else:
                form = form_model(initial=initial)
        else:
            form = form_model(initial=initial)
    return form


def get_index_items(**kwargs):
    """
    """
    context = {}
    app_settings_model = kwargs.get('app_settings_model')
    columns_visible = kwargs.get('columns_visible')
    company_model = kwargs.get('company_model')
    model = kwargs.get('model')
    order_by = kwargs.get('order_by')
    page_size = kwargs.get('page_size')
    request = kwargs.get('request')
    search_fields = kwargs.get('search_fields')
    show_search = kwargs.get('show_search')
    model_name = model._meta.verbose_name
    edit_url = '%s_edit' % model_name
    view_url = '%s_view' % model_name
    if columns_visible:
        context['columns_visible'] = columns_visible
    if company_model:
        company = company_model.get_solo()
        context['company'] = company
    page = get_query(request, 'page')
    paginated = get_query(request, 'paginated')
    search = get_query(request, 'search')
    if request:
        context['is_staff'] = request.user.is_staff
    # Search is easy
    if request.method == 'POST':
        if search == u'':  # Empty search returns none
            context['show_search'] = True
            return context
        else:
            return get_search_results(
                context,
                model,
                search_fields,
                search,
                app_settings_model=app_settings_model,
                edit_url=edit_url,
                view_url=view_url,
                request=request)
    # Not a search
    if model_name == 'time':
        items = model.objects.filter(estimate=None)
    elif model_name == 'invoice':
        items = model.objects.filter(last_payment_date=None)
    else:
        items = model.objects.all()
    if order_by is not None:  # Order items
        # http://stackoverflow.com/a/20257999/185820
        items = items.order_by(*order_by)
    if not request.user.is_authenticated:  # Don't show items to anon
        items = []
    if model_name == 'note':  # Per model extras
        context['note_info'] = get_note_info(model)
    elif model_name == 'time':
        context['total_hours'] = get_total_hours(items)
    if paginated:  # Paginate if paginated
        page_size = get_setting(
            request, app_settings_model, 'page_size', page_size=page_size)
        items = paginate(items, page, page_size)
    context['edit_url'] = edit_url
    context['view_url'] = view_url
    context['icon_size'] = get_setting(request, app_settings_model,
                                       'icon_size')
    context['icon_color'] = get_setting(request, app_settings_model,
                                        'icon_color')
    context['page'] = page
    context['paginated'] = paginated
    context['show_search'] = show_search
    items = set_items_name(model_name, items=items)
    context['items'] = items
    context['active_nav'] = model_name
    return context


def get_invoice_totals(invoices):
    invoice_amount = invoice_cog = 0
    for invoice in invoices:
        if invoice.amount:
            invoice_amount += float(invoice.amount)
        if invoice.cog:
            invoice_cog += float(invoice.cog)
    return invoice_amount, invoice_amount - invoice_cog


def get_note_info(note_model):
    note_info = {}
    active = len(note_model.objects.filter(active=True))
    inactive = len(note_model.objects.filter(active=False))
    hidden = len(note_model.objects.filter(hidden=True))
    not_hidden = inactive - hidden
    total = len(note_model.objects.all())
    note_info['active'] = active
    note_info['inactive'] = inactive
    note_info['hidden'] = hidden
    note_info['not_hidden'] = not_hidden
    note_info['total'] = total
    return note_info


def get_page_items(**kwargs):
    app_settings_model = kwargs.get('app_settings_model')
    company_model = kwargs.get('company_model')
    columns_visible = kwargs.get('columns_visible')
    contact_model = kwargs.get('contact_model')
    contract_model = kwargs.get('contract_model')
    dashboard_item_model = kwargs.get('dashboard_item_model')
    estimate_model = kwargs.get('estimate_model')
    invoice_model = kwargs.get('invoice_model')
    model = kwargs.get('model')
    note_model = kwargs.get('note_model')
    obj = kwargs.get('obj')
    project_model = kwargs.get('project_model')
    report_model = kwargs.get('report_model')
    request = kwargs.get('request')
    order_by = kwargs.get('order_by')
    pk = kwargs.get('pk')
    time_model = kwargs.get('time_model')
    user_model = kwargs.get('user_model')
    context = {}
    items = None
    if company_model:
        company = company_model.get_solo()
        context['company'] = company
    if columns_visible:
        context['columns_visible'] = columns_visible
    if model or obj:
        if model:
            model_name = model._meta.verbose_name
        elif obj:
            model_name = obj._meta.verbose_name
        context['model_name'] = model_name
        context['active_nav'] = model_name
        context['edit_url'] = '%s_edit' % model_name
        context['view_url'] = '%s_view' % model_name
        if model_name == 'Settings App':
            app_settings = app_settings_model.get_solo()
            context['items'] = get_fields([
                app_settings,
            ])  # table_items.html
        elif model_name == 'Settings Company':
            company_settings = model.get_solo()
            context['items'] = get_fields([
                company_settings,
            ])  # table_items.html
        elif model_name == 'Settings Contract':
            contract_settings = model.get_solo()
            context['items'] = get_fields([
                contract_settings,
            ])  # table_items.html
        elif model_name == 'client':
            client = get_object_or_404(model, pk=pk)
            contacts = contact_model.objects.filter(client=client)
            contracts = contract_model.objects.filter(client=client)
            projects = project_model.objects.filter(client=client)
            context['contacts'] = contacts
            context['contracts'] = contracts
            context['item'] = client
            context['notes'] = client.note.all()
            context['projects'] = projects
        elif model_name == 'contact':
            contact = get_object_or_404(model, pk=pk)
            context['items'] = get_fields([
                contact,
            ])  # table_items.html
            context['item'] = contact
        elif model_name == 'contract':
            contract = get_object_or_404(model, pk=pk)
            estimate = contract.statement_of_work
            times = None
            if estimate:
                times = time_model.objects.filter(
                    client=estimate.client,
                    estimate=None,
                    project=None,
                    invoiced=False,
                    invoice=None)
            context['doc_type'] = model_name
            context['item'] = contract
            context['times'] = times
        elif model_name == 'estimate':  # handle obj or model
            if not obj:
                estimate = get_object_or_404(model, pk=pk)
            else:
                estimate = obj
            if estimate.is_sow:
                doc_type = 'statement of work'
            elif estimate.is_to:
                doc_type = 'task order'
            else:
                doc_type = model_name
            times = time_model.objects.filter(estimate=estimate)
            if order_by:
                times = times.order_by(*order_by['time'])
            times = set_invoice_totals(times, estimate=estimate)
            total_hours = get_total_hours(times)
            context['doc_type'] = doc_type
            context['entries'] = times
            context['item'] = estimate
            context['total_hours'] = total_hours
        if model_name == 'file':
            file_obj = get_object_or_404(model, pk=pk)
            context['doc_type'] = model_name
            context['item'] = file_obj
        elif model_name == 'invoice':
            invoice = get_object_or_404(model, pk=pk)
            times = get_times_for_obj(invoice, time_model)
            times = times.order_by(*order_by['time'])
            times = set_invoice_totals(times, invoice=invoice)
            last_payment_date = invoice.last_payment_date
            total_hours = get_total_hours(times)
            context['doc_type'] = model_name
            context['entries'] = times
            context['item'] = invoice
            context['invoice'] = True
            context['last_payment_date'] = last_payment_date
            context['total_hours'] = total_hours
        elif model_name == 'newsletter':
            newsletter = get_object_or_404(model, pk=pk)
            context['doc_type'] = model_name
            context['item'] = newsletter
        elif model_name == 'note':
            note = get_object_or_404(model, pk=pk)
            context['item'] = note
        elif model_name == 'project':
            project = get_object_or_404(model, pk=pk)
            contacts = contact_model.objects.all()
            estimates = estimate_model.objects.filter(project=project)
            invoices = invoice_model.objects.filter(project=project)
            invoices = invoices.order_by('-issue_date')
            times = get_times_for_obj(project, time_model)
            times = times.order_by(*order_by['time'])
            users = user_model.objects.filter(project=project)
            items = set_items_name('contact', items=contacts)
            items = set_items_name('estimate', items=estimates, _items=items)
            items = set_items_name('invoice', items=invoices, _items=items)
            items = set_items_name('time', items=times, _items=items)
            items = set_items_name('user', items=users, _items=items)
            context['item'] = project
            context['items'] = items
        elif model_name == 'proposal':
            proposal = get_object_or_404(model, pk=pk)
            context['doc_type'] = model_name
            context['item'] = proposal
        elif model_name == 'report':
            report = get_object_or_404(model, pk=pk)
            reports = model.objects.filter(active=True)
            reports = reports.aggregate(
                gross=Sum(F('gross')), net=Sum(F('net')))
            context['cost'] = report.gross - report.net
            context['item'] = report
        elif model_name == 'task':
            task = get_object_or_404(model, pk=pk)
            context['item'] = task
        elif model_name == 'time':
            time_entry = get_object_or_404(model, pk=pk)
            context['item'] = time_entry
        elif model_name == 'user':
            exclude_fields = ('id', 'created', 'updated', 'hidden', 'active',
                              'app_admin', 'is_contact', 'notify', 'published',
                              'dashboard_override', 'dashboard_choices',
                              'editor', 'icon_size', 'icon_color', 'page_size',
                              'preferred_username', 'unit', 'avatar_url')
            user = get_object_or_404(model, pk=pk)
            projects = project_model.objects.filter(
                team__in=[
                    user,
                ], active=True)
            projects = projects.order_by(*order_by['project'])
            times = time_model.objects.filter(
                estimate=None, invoiced=False, user=user)
            times = times.order_by(*order_by['time'])
            contacts = contact_model.objects.all()
            context['item'] = user
            context['items'] = get_fields(
                [
                    user.profile,
                ], exclude_fields=exclude_fields)  # table_items.html
            context['projects'] = projects
            context['times'] = times
    else:  # home
        if request:
            if request.user.is_authenticated:
                # Dashboard
                dashboard_choices = get_setting(request, app_settings_model,
                                                'dashboard_choices')
                dashboard_items = [
                    i.title.lower()
                    for i in dashboard_item_model.objects.all()
                ]
                context['dashboard_choices'] = dashboard_choices
                context['dashboard_items'] = dashboard_items
                # Items
                invoices = invoice_model.objects.filter(last_payment_date=None)
                invoices = invoices.order_by(*order_by['invoice'])
                notes = note_model.objects.filter(active=True, hidden=False)
                notes = notes.order_by(*order_by['note'])
                projects = project_model.objects.filter(
                    active=True, hidden=False)
                projects = projects.order_by(*order_by['project'])
                times = time_model.objects.filter(
                    invoiced=False, user=request.user)
                times = times.order_by(*order_by['time'])
                items = set_items_name('invoice', items=invoices)
                items = set_items_name('note', items=notes, _items=items)
                items = set_items_name('project', items=projects, _items=items)
                items = set_items_name('time', items=times, _items=items)
                # Plot
                points = report_model.objects.filter(active=True)
                # Totals
                gross, net = get_invoice_totals(invoices)
                ip_address = request.META.get('HTTP_X_REAL_IP')
                context['gross'] = gross
                context['invoices'] = invoices
                context['geo_ip_data'] = get_geo_ip_data(request)
                context['ip_address'] = ip_address
                context['items'] = items
                context['net'] = net
                context['notes'] = notes
                context['note_info'] = get_note_info(note_model)
                context['points'] = points
                context['projects'] = projects
                context['times'] = times
                total_hours = get_total_hours(times)
                context['total_hours'] = total_hours
    if request:
        context['icon_size'] = get_setting(request, app_settings_model,
                                           'icon_size')
        context['icon_color'] = get_setting(request, app_settings_model,
                                            'icon_color')
        doc = get_query(request, 'doc')
        pdf = get_query(request, 'pdf')
        context['doc'] = doc
        context['pdf'] = pdf
        context['is_staff'] = request.user.is_staff
    return context


def get_plot(request):  # http://stackoverflow.com/a/5515994/185820
    """
    """
    points = get_query(request, 'points')
    # http://matplotlib.org/examples/api/date_demo.html
    x = [
        date2num(timezone.datetime.strptime(i[1], '%Y-%m-%d')) for i in points
    ]
    y = [i[0] for i in points]
    figure = Figure()
    canvas = FigureCanvasAgg(figure)
    axes = figure.add_subplot(1, 1, 1)
    axes.grid(True)
    axes.plot(x, y)
    axes.xaxis.set_major_formatter(DateFormatter('%m'))
    # write image data to a string buffer and get the PNG image bytes
    buf = BytesIO()
    canvas.print_png(buf)
    data = buf.getvalue()
    # write image bytes back to the browser
    return HttpResponse(data, content_type="image/png")


def get_query(request, query):
    """
    Special handling for some query strings
    """
    if query == 'paginated':
        paginated = request.GET.get('paginated')
        if paginated == u'false':
            return False
        else:
            return True
    elif query == 'search' and request.method == 'POST':
        return request.POST.get('search', '')
    elif query == 'points':  # plot
        points = request.GET.get('points')
        if points:
            points = points.split(' ')
        else:
            points = []
        points = [i.split(',') for i in points]
        return points
    elif query == 'checkbox':
        query_checkbox = {}
        query_checkbox_active = request.POST.get('checkbox-active')
        query_checkbox_subscribe = request.POST.get('checkbox-subscribe')
        condition = (  # if any of these exist
            query_checkbox_active == 'on' or query_checkbox_active == 'off'
            or query_checkbox_subscribe == 'on'
            or query_checkbox_subscribe == 'off')
        query_checkbox['active'] = query_checkbox_active
        query_checkbox['subscribe'] = query_checkbox_subscribe
        query_checkbox['condition'] = condition
        return query_checkbox
    elif query == 'doc':
        doc = request.GET.get('doc')
        if doc:
            return True
        else:
            return False
    elif query == 'pdf':
        pdf = request.GET.get('pdf')
        if pdf:
            return True
        else:
            return False
    elif query == 'test':
        test = request.GET.get('test')
        if test:
            return True
        else:
            return False
    else:  # Normal handling
        return request.GET.get(query, '')


def get_recipients(obj):
    """
    Returns first name and email address
    """
    model_name = obj._meta.verbose_name
    if model_name == 'contact':
        return [
            (obj.first_name, obj.email),
        ]
    elif model_name == 'estimate':
        return [(i.first_name, i.email) for i in obj.contacts.all()]
    elif model_name == 'note':
        return [(i.first_name, i.email) for i in obj.contacts.all()]
    elif model_name == 'time':
        return [('Alex', 'aclark@aclark.net')]


def get_search_results(context,
                       model,
                       search_fields,
                       search,
                       app_settings_model=None,
                       edit_url=None,
                       view_url=None,
                       request=None):
    query = []
    model_name = model._meta.verbose_name
    for field in search_fields:
        query.append(Q(**{field + '__icontains': search}))
    items = model.objects.filter(reduce(OR, query))
    context['active_nav'] = model_name
    context['edit_url'] = edit_url
    context['view_url'] = view_url
    context['icon_size'] = get_setting(request, app_settings_model,
                                       'icon_size')
    context['icon_color'] = get_setting(request, app_settings_model,
                                        'icon_color')
    context['show_search'] = True
    items = set_items_name(model_name, items=items)
    context['items'] = items
    return context


def get_setting(request, app_settings_model, setting, page_size=None):
    """
    Allow user to override global setting
    """
    if not request.user.is_authenticated:
        return
    dashboard_override = user_pref = None
    app_settings = app_settings_model.get_solo()
    if setting == 'icon_size':
        if has_profile(request.user):
            user_pref = request.user.profile.icon_size
        if user_pref:
            return user_pref
        else:
            return app_settings.icon_size
    elif setting == 'icon_color':
        if has_profile(request.user):
            user_pref = request.user.profile.icon_color
        if user_pref:
            return user_pref
        else:
            return app_settings.icon_color
    elif setting == 'page_size':
        if has_profile(request.user):
            user_pref = request.user.profile.page_size
        if user_pref:
            return user_pref
        elif page_size:  # View's page_size preference
            return page_size
        else:
            return app_settings.page_size
    elif setting == 'dashboard_choices':
        dashboard_choices = app_settings.dashboard_choices
        dashboard_override = False
        if has_profile(request.user):
            dashboard_override = request.user.profile.dashboard_override
        if has_profile(request.user) and dashboard_override:
            dashboard_choices = request.user.profile.dashboard_choices
        return dashboard_choices
    elif setting == 'exclude_hidden':
        return app_settings.exclude_hidden


def get_template_and_url(**kwargs):
    """
    """
    model_name = kwargs.get('model_name')
    page_type = kwargs.get('page_type')
    if page_type == 'view':
        url_name = URL_NAMES[model_name][0]
        template_name = '%s.html' % url_name
        return template_name, url_name
    elif page_type == 'copy':
        url_name = URL_NAMES[model_name][1]
        return url_name
    elif page_type == 'edit':
        template_name = '%s.html' % URL_NAMES[model_name][1]
        return template_name
    elif page_type == 'home':
        url_name = 'home'
        return url_name
    elif page_type == 'index':
        url_name = URL_NAMES[model_name][2]
        return url_name


def get_times_for_obj(obj, time_model):
    model_name = obj._meta.verbose_name
    if model_name == 'invoice':
        times = time_model.objects.filter(estimate=None, invoice=obj)
    elif model_name == 'project':
        times = time_model.objects.filter(
            invoiced=False, estimate=None, project=obj)
    return times


def get_total_earned(request, total_hours):
    total_earned = 0
    if has_profile(request.user):
        if request.user.profile.rate:
            rate = request.user.profile.rate
            total_earned = total_hours * rate
    return '%.2f' % total_earned


def get_total_hours(items):
    total_hours = items.aggregate(hours=Sum(F('hours')))
    total_hours = total_hours['hours']
    return total_hours


def gravatar_url(email):
    """
    MD5 hash of email address for use with Gravatar. Return generic
    if none exists.
    """
    try:
        return django_settings.GRAVATAR_URL % md5(email.lower()).hexdigest()
    except:
        # https://stackoverflow.com/a/7585378/185820
        return django_settings.GRAVATAR_URL % md5(
            'db@aclark.net'.encode('utf-8')).hexdigest()


def has_profile(user):
    return hasattr(user, 'profile')


def last_month():
    """
    Returns last day of last month
    """
    first = timezone.now().replace(day=1)
    return first - timezone.timedelta(days=1)


def mail_compose(obj, **kwargs):
    """
    Compose message based on object type.
    """
    model_name = obj._meta.verbose_name
    # Get kwargs
    form = kwargs.get('form')
    hostname = kwargs.get('hostname')
    mail_from = kwargs.get('mail_from')
    mail_to = kwargs.get('mail_to')
    status_message = kwargs.get('status_message')
    time_model = kwargs.get('time_model')
    # Conditionally create message
    if model_name == 'contact':
        message = form.cleaned_data['message']
        subject = form.cleaned_data['subject']
    elif model_name == 'estimate':
        message = render_to_string('pdf_invoice.html',
                                   get_page_items(
                                       obj=obj, time_model=time_model))
        subject = obj.subject
    elif model_name == 'note':
        message = obj.note
        subject = obj.title
    elif model_name == 'time':
        message = '%s' % obj.get_absolute_url(hostname)
        subject = status_message.get('success')

    # first_name = kwargs.get('first_name')
    # if first_name:
    #     message = render_to_string('first_name.html', {
    #         'first_name': first_name,
    #         'message': message,
    #     })

    # If a form is passed in, use the mail template selected by the user.
    context = {}
    if form:  # http://stackoverflow.com/a/28476681/185820
        if 'send_html' in form.data:
            html_message = render_to_string(form.data['template'], {
                'message': message,
            })
            context['html_message'] = html_message
    else:  # No form, must be someone running `python manage.py send_note`.
        context['html_message'] = render_to_string('html_message.html', {
            'message': message,
        })
    context['mail_from'] = mail_from
    context['mail_to'] = mail_to
    context['message'] = message
    context['subject'] = subject
    return context


def mail_obj(request, **kwargs):
    query_contact = get_query(request, 'contact')
    query_estimate = get_query(request, 'estimate')
    query_note = get_query(request, 'note')
    contact_model = kwargs.get('contact_model')
    estimate_model = kwargs.get('estimate_model')
    note_model = kwargs.get('note_model')
    if contact_model and query_contact:
        obj = contact_model.objects.get(pk=query_contact)
    elif note_model and query_note:
        obj = note_model.objects.get(pk=query_note)
    elif estimate_model and query_estimate:
        obj = estimate_model.objects.get(pk=query_estimate)
    return obj


def mail_process(obj, **kwargs):
    """
    """
    form = kwargs.get('form')
    request = kwargs.get('request')
    status_message = kwargs.get('status_message')
    time_model = kwargs.get('time_model')
    # Iterate over recipients, compose and send mail to
    # each.
    hostname = request.META.get('HTTP_HOST')
    mail_from = django_settings.EMAIL_FROM
    recipients = get_recipients(obj)
    for first_name, email_address in recipients:
        status = mail_send(**mail_compose(
            obj,
            form=form,
            first_name=first_name,
            hostname=hostname,
            mail_from=mail_from,
            mail_to=email_address,
            status_message=status_message,
            time_model=time_model,
            request=request))
    if status:
        messages.add_message(request, messages.SUCCESS,
                             status_message.get('success') % recipients)
    else:
        messages.add_message(request, messages.WARNING,
                             status_message.get('failure') % recipients)


def mail_send(**kwargs):
    html_message = kwargs.get('html_message')
    mail_from = kwargs.get('mail_from')
    mail_to = kwargs.get('mail_to')
    message = kwargs.get('message')
    subject = kwargs.get('subject')
    try:
        send_mail(
            subject,
            message,
            mail_from, (mail_to, ),
            fail_silently=False,
            html_message=html_message)
        status = True
    except BotoServerError:
        status = False
    return status


def obj_copy(obj):
    dup = obj
    dup.pk = None
    dup.save()
    kwargs = {}
    kwargs['pk'] = dup.pk
    model_name = obj._meta.verbose_name
    url_name = get_template_and_url(model_name=model_name, page_type='copy')
    return HttpResponseRedirect(reverse(url_name, kwargs=kwargs))


def obj_redir(obj, pk=None):
    """
    Redir after edit, special cases for some objects
    """
    model_name = obj._meta.verbose_name
    template_name, url_name = get_template_and_url(
        model_name=model_name, page_type='view')  # Redir to view
    kwargs = {}
    if pk:  # Exists
        kwargs['pk'] = pk
        if model_name == 'Settings App':  # Special cases for settings
            return HttpResponseRedirect(reverse(url_name))
        elif model_name == 'Settings Company':
            return HttpResponseRedirect(reverse(url_name))
        elif model_name == 'Settings Contract':
            return HttpResponseRedirect(reverse(url_name))
    else:  # New
        if model_name == 'profile':  # One of to create profile for new
            kwargs['pk'] = obj.user.pk  # user
        else:
            kwargs['pk'] = obj.pk
    return HttpResponseRedirect(reverse(url_name, kwargs=kwargs))


def obj_remove(obj):
    model_name = obj._meta.verbose_name
    if model_name == 'time':
        url_name = get_template_and_url(
            model_name=model_name, page_type='home')  # Redir to home
    else:
        url_name = get_template_and_url(
            model_name=model_name, page_type='index')  # Redir to index
    if model_name == 'profile':
        obj.user.delete()
    else:
        obj.delete()
    return HttpResponseRedirect(reverse(url_name))


def obj_sent(obj, ref):
    """
    Mark time entry as invoiced when invoice sent.
    """
    for time in obj.time_set.all():
        time.invoiced = True
        time.save()
    return HttpResponseRedirect(ref)


def paginate(items, page, page_size):
    """
    """
    paginator = Paginator(items, page_size, orphans=5)
    try:
        items = paginator.page(page)
    except PageNotAnInteger:
        items = paginator.page(1)
    except EmptyPage:
        items = paginator.page(paginator.num_pages)
    return items


def set_check_boxes(obj, query_checkbox, ref, app_settings_model):
    model_name = obj._meta.verbose_name
    if (query_checkbox['active'] == 'on'
            or query_checkbox['active'] == 'off'):  # Active
        if query_checkbox['active'] == 'on':
            obj.active = True
            obj.hidden = False
        else:
            obj.active = False
            if model_name == 'note' and app_settings_model:
                app_settings = app_settings_model.get_solo()
                if app_settings.auto_hide:  # Auto-hide
                    obj.hidden = True
    elif (query_checkbox['subscribe'] == 'on'
          or query_checkbox['subscribe'] == 'off'):  # Subscribe
        if query_checkbox['active'] == 'on':
            obj.subscribed = True
        else:
            obj.subscribed = False
    obj.save()
    return HttpResponseRedirect(ref)


def set_invoice_totals(times, estimate=None, invoice=None):
    """
    Set invoice, estimate and time totals
    """
    invoice_amount = invoice_cog = 0
    time_entry_amount = time_entry_cog = 0
    for time_entry in times:
        hours = time_entry.hours
        if time_entry.task:
            rate = time_entry.task.rate
            time_entry_amount = rate * hours
        if time_entry.user:
            rate = time_entry.user.profile.rate
            if rate:
                time_entry_cog = rate * hours
        time_entry.amount = '%.2f' % time_entry_amount
        time_entry.cog = '%.2f' % time_entry_cog
        invoice_amount += time_entry_amount
        invoice_cog += time_entry_cog
    if invoice:
        invoice.amount = '%.2f' % invoice_amount
        invoice.cog = '%.2f' % invoice_cog
        invoice.save()
    elif estimate:
        estimate.amount = '%.2f' % invoice_amount
        estimate.save()
    return times


def set_items_name(model_name, items=None, _items={}):
    """
    Share templates by returning dictionary of items e.g.
        for item in items.reports
    instead of:
        for item in reports
    """
    _items['%ss' % model_name] = items
    return _items


def set_relationship(obj, request, **kwargs):
    """
    Sets relationships after create or edit
    """
    client_model = kwargs.get('client_model')
    company_model = kwargs.get('company_model')
    estimate_model = kwargs.get('estimate_model')
    invoice_model = kwargs.get('invoice_model')
    project_model = kwargs.get('project_model')
    model_name = obj._meta.verbose_name
    if model_name == 'contact':
        query_client = get_query(request, 'client')
        if query_client:
            client = get_object_or_404(client_model, pk=query_client)
            obj.client = client
            obj.save()
    elif model_name == 'estimate' or model_name == 'invoice':
        query_project = get_query(request, 'project')
        if query_project:
            project = get_object_or_404(project_model, pk=query_project)
            obj.client = project.client
            obj.project = project
            obj.save()
    elif model_name == 'note':
        query_client = get_query(request, 'client')
        query_company = get_query(request, 'company')
        if query_client:
            client = get_object_or_404(client_model, pk=query_client)
            client.note.add(obj)
            client.save()
        elif query_company:
            company = company_model.get_solo()
            company.note.add(obj)
            company.save()
    elif model_name == 'project':
        query_client = get_query(request, 'client')
        if query_client:
            client = get_object_or_404(client_model, pk=query_client)
            obj.client = client
            obj.save()
    elif model_name == 'time':
        obj.user = request.user
        query_estimate = get_query(request, 'estimate')
        query_invoice = get_query(request, 'invoice')
        query_project = get_query(request, 'project')
        if not request.user.is_staff:  # Staff have more than one project
            user_projects = project_model.objects.filter(
                team__in=[
                    obj.user,
                ])
            if len(user_projects) > 0:
                obj.project = user_projects[0]
                obj.task = obj.project.task
        if query_estimate:
            estimate = get_object_or_404(estimate_model, pk=query_estimate)
            obj.estimate = estimate
        if query_invoice:
            invoice = get_object_or_404(invoice_model, pk=query_invoice)
            obj.invoice = invoice
            obj.task = invoice.project.task
        if query_project:
            project = get_object_or_404(project_model, pk=query_project)
            obj.project = project
            obj.save()  # Need save her to set task
            obj.task = project.task
        obj.save()
