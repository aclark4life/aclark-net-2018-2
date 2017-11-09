from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import StringIO
from lxml import etree


def generate_doc(contract):
    """
    https://stackoverflow.com/a/24122313/185820
    """
    document = Document()
    # Head
    task = ''
    if contract.task:
        task = contract.task
    title = document.add_heading(
        'ACLARK.NET, LLC %s AGREEMENT PREPARED FOR:' % task, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if contract.client:
        client_name = document.add_heading(contract.client.name, level=1)
        client_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_address = document.add_heading(contract.client.address, level=1)
        client_address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parser = etree.HTMLParser()  # http://lxml.de/parsing.html
    tree = etree.parse(StringIO(contract.body), parser)
    # Body
    for element in tree.iter():
        if element.tag == 'h2':
            document.add_heading(element.text, level=2)
        elif element.tag == 'p':
            document.add_paragraph(element.text)
    return document
